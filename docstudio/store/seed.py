"""First-run workspace seeding: 3 doc-type templates, 1 word template, and
1 example document with drafted chapters, sources, an open question, and 2
saved versions — so the demo works immediately after `python app.py`.
Idempotent: skipped if any doc-type templates already exist.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument

from docstudio import ingest
from docstudio.engine.mock import _ACRONYM_DEFS
from docstudio.formatter.templated import STYLES_END_MARKER, STYLES_START_MARKER
from docstudio.models import ChapterFrontmatter, ChapterSpec, DocTypeTemplate, InterviewQuestion, SourceRef
from docstudio.store.templates import render_doc_type_template


def seed_workspace(state) -> None:
    templates = state.templates
    if any(templates.workspace.doc_type_templates.glob("*.md")):
        return  # already seeded

    _seed_word_template(templates)
    _seed_doc_type_templates(templates)
    _seed_example_document(state)


# ---------------------------------------------------------------------------
# Doc-type templates
# ---------------------------------------------------------------------------

def _seed_doc_type_templates(templates) -> None:
    for tpl in (_fsd_template(), _integration_spec_template(), _approach_doc_template()):
        templates.save_doc_type(tpl)


def _fsd_template() -> DocTypeTemplate:
    return DocTypeTemplate(
        doc_type="fsd",
        name="Functional Specification Document",
        version="1.0",
        word_template="corporate-default.docx",
        general_instructions=(
            "A Functional Specification Document (FSD) is the artifact a client's architects, business "
            "stakeholders, and delivery team sign off on before build starts. It must give a build team "
            "enough detail to estimate and implement without further discovery, and give a business "
            "stakeholder enough clarity to approve scope.\n\n"
            "Tone: formal, third person, no marketing language. Every functional requirement gets a unique "
            "ID (FR-nn) and every non-functional requirement a measurable target — no vague terms like "
            "\"fast\" or \"secure\" without a number attached. Acronyms and domain terms are defined once, "
            "in the Glossary, and used consistently everywhere else."
        ),
        clarification_policy=(
            "Ask only when the sources don't already answer the question and the answer would materially "
            "change what gets written — not for things you can reasonably assume or infer. Prefer "
            "multiple-choice questions when the space of plausible answers is small (e.g. which payment "
            "rails, which environment); use free text only when it genuinely can't be enumerated. Batch up "
            "to 3 questions together if a chapter needs more than one thing clarified, rather than pausing "
            "repeatedly. Scope, in particular, blocks almost everything downstream — resolve it early."
        ),
        chapters=[
            ChapterSpec(
                number="01",
                title="Introduction",
                required=True,
                prompt=(
                    "This chapter must answer: what system, what business context, who are the "
                    "stakeholders.\nTone: formal, third person. Length: 300-500 words."
                ),
            ),
            ChapterSpec(
                number="02",
                title="Scope",
                required=True,
                prompt="In-scope and out-of-scope as two lists. Must be unambiguous.",
            ),
            ChapterSpec(
                number="03",
                title="Functional Requirements",
                required=True,
                prompt=(
                    "Describe each functional requirement with an ID, description, and priority.\n"
                    "Group by business capability. Reference source material where possible."
                ),
            ),
            ChapterSpec(
                number="04",
                title="Non-Functional Requirements",
                required=True,
                prompt="Performance, availability, security, and compliance requirements. Be specific and measurable.",
            ),
            ChapterSpec(
                number="05",
                title="Assumptions and Dependencies",
                required=True,
                prompt="List assumptions made and dependencies on other teams, systems, or third parties.",
            ),
            ChapterSpec(
                number="90",
                title="Glossary",
                required=True,
                derived=True,
                prompt=(
                    "Extract all domain terms, acronyms, and system names used across all chapters.\n"
                    "One table: Term | Definition."
                ),
            ),
        ],
        interview_bank=[
            InterviewQuestion(
                q="What is the target go-live date?",
                chapter="01",
                choices=[],
                context="The Introduction chapter frames the engagement timeline for readers, and downstream chapters (timeline, dependencies) will reference this date.",
            ),
            InterviewQuestion(
                q="Which payment rails are in scope?",
                chapter="02",
                choices=["FAST", "MEPS+", "SWIFT", "Multiple / other"],
                context="This is the single biggest driver of scope — it determines which integration chapters are even needed.",
            ),
            InterviewQuestion(
                q="Is card scheme integration in scope for this phase?",
                chapter="02",
                choices=["Yes", "No", "Deferred to phase 2"],
            ),
        ],
        quality_checklist=[
            "No chapter contradicts the scope chapter",
            "All acronyms appear in the glossary",
            "Every functional requirement has a unique ID",
        ],
    )


def _integration_spec_template() -> DocTypeTemplate:
    return DocTypeTemplate(
        doc_type="integration-spec",
        name="Integration Specification",
        version="1.0",
        word_template="corporate-default.docx",
        general_instructions=(
            "An Integration Specification is read by the engineers building both sides of an integration "
            "(the client's team and the delivery team), plus whoever signs off on non-functional targets. "
            "It must be precise enough that two teams who never talk directly can build compatible "
            "implementations from it alone.\n\n"
            "Tone: technical, precise, third person. Every message field gets an explicit type and "
            "required/optional flag — no implied structure. Every error code gets a documented retry "
            "behaviour. Prefer tables and Mermaid sequence diagrams over prose wherever a flow or "
            "structure can be shown instead of described."
        ),
        clarification_policy=(
            "Ask only when the sources don't already answer the question and the answer would materially "
            "change the design — not for things inferable from context. Prefer multiple-choice for "
            "questions with a small answer space (e.g. sync vs async). Batch up to 3 questions per chapter "
            "rather than pausing repeatedly. Sync/async and peak-volume answers block message-format and "
            "SLA chapters downstream, so resolve those early."
        ),
        chapters=[
            ChapterSpec(
                number="01",
                title="Integration Overview",
                required=True,
                prompt="Describe the systems being integrated, the integration pattern (sync/async, batch/real-time), and the business trigger for each flow.",
            ),
            ChapterSpec(
                number="02",
                title="Message Formats",
                required=True,
                prompt="Define request/response payloads per message type. Include field-level detail as a table: Field | Type | Required | Notes.",
            ),
            ChapterSpec(
                number="03",
                title="Sequence Flows",
                required=True,
                prompt="Describe the end-to-end sequence for each integration flow. Use a Mermaid sequence diagram where useful.",
            ),
            ChapterSpec(
                number="04",
                title="Error Handling and Retries",
                required=True,
                prompt="Enumerate error codes, retry policy, idempotency guarantees, and dead-letter handling.",
            ),
            ChapterSpec(
                number="05",
                title="SLAs and Monitoring",
                required=True,
                prompt="Define latency and availability targets, and how the integration will be monitored/alerted.",
            ),
            ChapterSpec(
                number="90",
                title="Glossary",
                required=True,
                derived=True,
                prompt="Extract all domain terms, acronyms, and system names used across all chapters.\nOne table: Term | Definition.",
            ),
        ],
        interview_bank=[
            InterviewQuestion(q="Is the integration synchronous or asynchronous?", chapter="01", choices=["Synchronous", "Asynchronous", "Both"]),
            InterviewQuestion(q="What is the expected peak transactions per second?", chapter="05", choices=[]),
        ],
        quality_checklist=[
            "Every message field has a defined type",
            "Every error code has a documented retry behaviour",
            "All acronyms appear in the glossary",
        ],
    )


def _approach_doc_template() -> DocTypeTemplate:
    return DocTypeTemplate(
        doc_type="approach-doc",
        name="Approach Document",
        version="1.0",
        word_template="corporate-default.docx",
        general_instructions=(
            "An Approach Document is a pre-sales / early-engagement artifact read by senior client "
            "stakeholders deciding whether to proceed, plus the delivery team who will be held to what it "
            "promises. It sells the approach while committing to something realistic.\n\n"
            "Tone: confident but not salesy, plain language over jargon — a senior stakeholder skimming "
            "the Executive Summary alone should understand the engagement. Every risk gets a named owner "
            "and a concrete mitigation, not a generic one. Timeline milestones must align with the phases "
            "named in Methodology."
        ),
        clarification_policy=(
            "Ask only when the sources don't already answer the question and the answer would materially "
            "change the proposed approach — not for things inferable from context. Prefer multiple-choice "
            "for questions with a small answer space (e.g. pricing model). Batch up to 3 questions per "
            "chapter rather than pausing repeatedly. Engagement start date and pricing model block the "
            "Timeline and Team/Governance chapters, so resolve those early."
        ),
        chapters=[
            ChapterSpec(
                number="01",
                title="Executive Summary",
                required=True,
                prompt="Summarize the engagement objective, approach, and expected outcome in plain language for a senior stakeholder. Length: 200-300 words.",
            ),
            ChapterSpec(
                number="02",
                title="Objectives and Success Criteria",
                required=True,
                prompt="List measurable objectives and how success will be evaluated.",
            ),
            ChapterSpec(
                number="03",
                title="Methodology",
                required=True,
                prompt="Describe the delivery methodology, phases, and key activities per phase.",
            ),
            ChapterSpec(
                number="04",
                title="Timeline and Milestones",
                required=True,
                prompt="Provide a phased timeline with key milestones. A Mermaid gantt or table is acceptable.",
            ),
            ChapterSpec(
                number="05",
                title="Team and Governance",
                required=True,
                prompt="Describe the proposed team structure, roles, and governance/escalation model.",
            ),
            ChapterSpec(
                number="06",
                title="Risks and Mitigations",
                required=True,
                prompt="List key risks with likelihood, impact, and mitigation as a table.",
            ),
            ChapterSpec(
                number="90",
                title="Glossary",
                required=True,
                derived=True,
                prompt="Extract all domain terms, acronyms, and system names used across all chapters.\nOne table: Term | Definition.",
            ),
        ],
        interview_bank=[
            InterviewQuestion(q="What is the target engagement start date?", chapter="04", choices=[]),
            InterviewQuestion(q="Is this a fixed-price or time-and-materials engagement?", chapter="05", choices=["Fixed-price", "Time and materials", "Hybrid"]),
        ],
        quality_checklist=[
            "Every risk has a named owner and mitigation",
            "Timeline milestones align with the methodology phases",
            "All acronyms appear in the glossary",
        ],
    )


# ---------------------------------------------------------------------------
# Word template
# ---------------------------------------------------------------------------

def _seed_word_template(templates) -> None:
    """A minimal but genuinely functional example of the {VARIABLE} /
    {CHAPTERS} / ##STYLES## convention TemplatedDocFormatter understands —
    not just descriptive prose. Attach a real corporate .docx from the
    Templates tab to replace it; this just makes the out-of-the-box export
    demonstrate the real mechanism instead of doing nothing.
    """
    doc = DocxDocument()
    doc.add_heading("{DOCUMENT_TITLE}", level=0)
    doc.add_paragraph("Prepared for: {CLIENT_NAME}")
    doc.add_paragraph("Author: {AUTHOR_NAME}")
    doc.add_paragraph(
        "This is the seeded placeholder corporate Word template — attach your "
        "own .docx from the Templates tab to replace it."
    )
    doc.add_heading("Document Chapters", level=1)
    doc.add_paragraph("{CHAPTERS}")

    doc.add_paragraph(STYLES_START_MARKER)
    doc.add_heading("HEADING", level=1)
    doc.add_heading("HEADING2", level=2)
    doc.add_heading("HEADING3", level=3)
    doc.add_paragraph("Normal")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Header"
    table.rows[0].cells[1].text = "Header"
    table.rows[1].cells[0].text = "Cell"
    table.rows[1].cells[1].text = "Cell"
    doc.add_paragraph(STYLES_END_MARKER)

    dest = templates.workspace.word_templates / "corporate-default.docx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))


# ---------------------------------------------------------------------------
# Example document
# ---------------------------------------------------------------------------

_INTRO_BODY = """# Introduction

Meridian Bank has engaged the consulting team to design and specify a new Payment
Gateway that will sit in front of the bank's core banking platform and
provide a single, channel-agnostic entry point for retail and corporate
payment initiation. The current landscape relies on a set of point-to-point
integrations between individual channels (internet banking, mobile, branch)
and the payment rails (FAST, MEPS+, SWIFT), which has become costly to
maintain and slow to extend to new rails or new channels.

The Payment Gateway will provide a unified API surface for payment
initiation, status enquiry, and reconciliation, decoupling channel teams from
the specifics of each payment rail. This document specifies the functional
and non-functional requirements for the first phase of the Payment Gateway,
covering retail FAST payments and corporate MEPS+ payments.

**Key stakeholders:**

- Treasury Operations
- the Client Integration squad
- Risk & Fraud

This specification will be used as the basis for solution design and UAT
planning, and reflects the outcomes of the discovery workshops held with
Meridian Bank Treasury Operations and the Client Integration squad in July 2026.
"""

_SCOPE_BODY = """# Scope

## In Scope

- Payment initiation via the new gateway for retail (FAST) and corporate (MEPS+) channels
- Real-time transaction status notifications to the client's core system
- Reconciliation and exception reporting for settled transactions

## Out of Scope

- Card issuing and card management services
- Historical data migration beyond the last 12 months
- Branch teller integration (handled by a separate workstream)

Card scheme integration for this phase is still being confirmed with the
client's Risk & Fraud team.

> [OPEN QUESTION: Is card scheme integration in scope for this phase?]
"""

_FUNCTIONAL_BODY = """# Functional Requirements

| ID | Requirement | Priority |
|---|---|---|
| FR-01 | The gateway must accept a payment initiation request via a REST API and validate it against the ISO 20022 schema before routing. | Must |
| FR-02 | The gateway must route FAST payments to the FAST rail and MEPS+ payments to the MEPS+ rail based on the payment type field. | Must |
| FR-03 | The gateway must expose a status enquiry endpoint returning the current lifecycle state of a payment (initiated, processing, settled, rejected). | Must |
| FR-04 | The gateway must produce a T+1 reconciliation file for Treasury Operations, flagging any unmatched transactions as exceptions. | Should |

Straight-through processing is required end-to-end for FR-01 and FR-02, with
manual intervention reserved for exception handling in FR-04.
"""

_ASSUMPTIONS_BODY = """> _Not yet drafted._
"""


def _seed_example_document(state) -> None:
    documents = state.documents
    manifest = documents.create_document(
        title="Meridian Bank Payment Gateway FSD", doc_type="fsd", client="Meridian Bank"
    )
    slug = manifest.slug

    _write_chapter(documents, slug, "01-introduction.md", _INTRO_BODY, status="final")
    _write_chapter(documents, slug, "02-scope.md", _SCOPE_BODY, status="draft", open_questions=1)
    _write_chapter(documents, slug, "03-functional-requirements.md", _FUNCTIONAL_BODY, status="draft")
    documents.append_decision(
        slug,
        "Is card scheme integration in scope for this phase?",
        "(deferred)",
        chapter="02-scope.md",
    )

    # Sources: one text source, one embed-mode spreadsheet, ingested through
    # the real pipeline so "view extracted" works exactly as it would for an
    # uploaded file.
    _seed_text_source(documents, slug)
    _seed_spreadsheet_source(documents, slug)

    # Glossary refresh, hand-assembled from the drafted chapters above.
    glossary_body = _glossary_for_seed()
    _write_chapter(documents, slug, "90-glossary.md", glossary_body, status="draft")

    manifest = documents.get_manifest(slug)
    manifest.status = "drafting"
    documents.save_manifest(manifest)

    # v1
    state.versions.save_version(slug)

    # A small edit so v1 -> current has a meaningful diff, then v2.
    chapter = documents.get_chapter(slug, "03-functional-requirements.md")
    documents.save_chapter(
        slug,
        "03-functional-requirements.md",
        chapter.body + "\n\nFR-04 reconciliation exceptions must be surfaced in the Treasury Operations dashboard within 15 minutes of file generation.\n",
        chapter.frontmatter,
    )
    state.versions.save_version(slug)


def _write_chapter(documents, slug: str, file: str, body: str, status: str, open_questions: int = 0) -> None:
    chapter = documents.get_chapter(slug, file)
    frontmatter = ChapterFrontmatter(
        title=chapter.frontmatter.title, status=status, sources_used=[], last_generated=datetime.now()
    )
    documents.save_chapter(slug, file, body, frontmatter)
    if open_questions:
        documents.set_open_questions(slug, file, open_questions)


def _seed_text_source(documents, slug: str) -> None:
    filename = "client-email-thread.txt"
    content = (
        "From: treasury.ops@meridianbank-example.com\n"
        "Subject: RE: Payment Gateway kickoff - scope confirmation\n\n"
        "Confirming from our side: phase 1 should cover FAST for retail and "
        "MEPS+ for corporate payments. SWIFT is out of scope for now, revisit "
        "in phase 2. We still need Risk & Fraud to confirm whether card scheme "
        "integration is in or out for phase 1 - please chase.\n\n"
        "Target go-live is end of Q1 2027, T+1 reconciliation is a hard "
        "requirement from Treasury Operations.\n"
    )
    original = documents.originals_dir(slug) / filename
    original.parent.mkdir(parents=True, exist_ok=True)
    original.write_text(content, encoding="utf-8")

    result = ingest.extract(original)
    extracted_name = f"{Path(filename).stem}.md"
    (documents.extracted_dir(slug)).mkdir(parents=True, exist_ok=True)
    (documents.extracted_dir(slug) / extracted_name).write_text(result.markdown or "", encoding="utf-8")

    documents.add_source_ref(
        slug,
        SourceRef(
            id=documents.new_source_id(),
            file=f"originals/{filename}",
            extracted=f"extracted/{extracted_name}",
            label="Client email thread",
            mode="source",
            extraction_status=result.status,  # type: ignore[arg-type]
            content_type="text/plain",
        ),
    )


def _seed_spreadsheet_source(documents, slug: str) -> None:
    from openpyxl import Workbook

    filename = "rail-fee-schedule.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Fees"
    ws.append(["Rail", "Fee Type", "Amount (SGD)"])
    ws.append(["FAST", "Per-transaction", 0.20])
    ws.append(["MEPS+", "Per-transaction", 5.00])
    ws.append(["SWIFT", "Per-transaction", 15.00])
    original = documents.originals_dir(slug) / filename
    original.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(original))

    result = ingest.extract(original)
    extracted_name = f"{Path(filename).stem}.md"
    (documents.extracted_dir(slug) / extracted_name).write_text(result.markdown or "", encoding="utf-8")

    documents.add_source_ref(
        slug,
        SourceRef(
            id=documents.new_source_id(),
            file=f"originals/{filename}",
            extracted=f"extracted/{extracted_name}",
            label="Rail fee schedule",
            mode="embed",
            extraction_status=result.status,  # type: ignore[arg-type]
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
    )


def _glossary_for_seed() -> str:
    terms = ["FAST", "MEPS+", "SWIFT", "API", "ISO 20022", "STP", "T+1"]
    rows = ["| Term | Definition |", "|---|---|"]
    for t in terms:
        rows.append(f"| {t} | {_ACRONYM_DEFS.get(t, 'Domain term referenced across the document chapters.')} |")
    return "# Glossary\n\n" + "\n".join(rows) + "\n"
