"""Parses templates/doc_types/*.md into DocTypeTemplate models, and manages
templates/word_templates/*.docx.
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

import yaml

from docstudio.models import ChapterSpec, DocTypeTemplate, InterviewQuestion
from docstudio.settings import WorkspacePaths

_FRONTMATTER_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n(.*)$", re.DOTALL)


def _split_frontmatter(text: str) -> tuple[dict, str]:
    m = _FRONTMATTER_RE.match(text)
    if not m:
        return {}, text
    meta = yaml.safe_load(m.group(1)) or {}
    return meta, m.group(2)


def _split_sections(text: str, prefix: str) -> list[tuple[str, str]]:
    pattern = re.compile(rf"^{re.escape(prefix)} (.+?)\s*$", re.MULTILINE)
    matches = list(pattern.finditer(text))
    out = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        out.append((m.group(1).strip(), text[start:end]))
    return out


def parse_doc_type_template(text: str) -> DocTypeTemplate:
    meta, body = _split_frontmatter(text)
    top = dict(_split_sections(body, "#"))

    chapters: list[ChapterSpec] = []
    for header, block in _split_sections(top.get("Chapters", ""), "##"):
        number, _, title = header.partition(" - ")
        data = yaml.safe_load(block) or {}
        chapters.append(
            ChapterSpec(
                number=number.strip(),
                title=title.strip(),
                required=bool(data.get("required", True)),
                derived=bool(data.get("derived", False)),
                prompt=(data.get("prompt") or "").strip(),
            )
        )
    chapters.sort(key=lambda c: c.number)

    interview_raw = yaml.safe_load(top.get("Interview Bank", "")) or []
    interview_bank = [InterviewQuestion(**q) for q in interview_raw]

    checklist = yaml.safe_load(top.get("Quality Checklist", "")) or []

    return DocTypeTemplate(
        doc_type=meta.get("doc_type", ""),
        name=meta.get("name", meta.get("doc_type", "")),
        version=str(meta.get("version", "1.0")),
        word_template=meta.get("word_template", "") or "",
        chapters=chapters,
        interview_bank=interview_bank,
        quality_checklist=list(checklist),
    )


def render_doc_type_template(tpl: DocTypeTemplate) -> str:
    """Inverse of parse_doc_type_template. Every scalar goes through
    yaml.safe_dump so titles/prompts/questions containing quotes, colons, or
    newlines round-trip correctly (this format is user-editable via the
    Templates UI, not just hand-written).
    """
    frontmatter = {"doc_type": tpl.doc_type, "name": tpl.name, "version": tpl.version}
    if tpl.word_template:
        frontmatter["word_template"] = tpl.word_template
    lines = ["---", yaml.safe_dump(frontmatter, sort_keys=False, allow_unicode=True).strip(), "---", "", "# Chapters", ""]

    for c in tpl.chapters:
        lines.append(f"## {c.number} - {c.title}")
        block = {"required": c.required}
        if c.derived:
            block["derived"] = True
        block["prompt"] = c.prompt
        lines.append(yaml.safe_dump(block, sort_keys=False, allow_unicode=True).strip())
        lines.append("")

    lines.append("# Interview Bank")
    lines.append("")
    if tpl.interview_bank:
        bank = [
            {"q": q.q, "chapter": q.chapter, "choices": list(q.choices), "context": q.context}
            for q in tpl.interview_bank
        ]
        lines.append(yaml.safe_dump(bank, sort_keys=False, allow_unicode=True).strip())
    lines.append("")

    lines.append("# Quality Checklist")
    lines.append("")
    if tpl.quality_checklist:
        lines.append(yaml.safe_dump(list(tpl.quality_checklist), sort_keys=False, allow_unicode=True).strip())
    lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Word template {VARIABLE} scanning
# ---------------------------------------------------------------------------

# ALL_CAPS convention only — excludes decorative/lowercase tokens like the
# copyright-year "{yyyy}" some corporate footers carry, which formatter code
# substitutes automatically rather than surfacing as a document variable.
VARIABLE_PATTERN = re.compile(r"\{\s*([A-Z][A-Z0-9_]*)\s*\}")
CHAPTERS_MARKER = "CHAPTERS"


def _iter_paragraphs(container):
    """Recursively yield every paragraph in a document/cell/header/footer,
    including nested tables inside table cells.
    """
    yield from container.paragraphs
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _iter_containers(document):
    """All paragraph-bearing parts of a document: the body plus every
    section's headers/footers (first-page and even-page variants included).
    """
    yield document
    for section in document.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                part.paragraphs  # noqa: B018 - probes whether the part is usable
            except Exception:
                continue
            yield part


def scan_template_variables(docx_path: Path) -> list[str]:
    """Distinct {VARIABLE} tokens used anywhere in a .docx (body, tables,
    nested tables, headers, footers) — used both to build the "fill these in"
    form and, symmetrically, at export time to know what to substitute.
    """
    if not docx_path.exists():
        return []
    from docx import Document as DocxDocument

    doc = DocxDocument(str(docx_path))
    found: set[str] = set()
    for container in _iter_containers(doc):
        for p in _iter_paragraphs(container):
            for m in VARIABLE_PATTERN.finditer(p.text):
                found.add(m.group(1))
    found.discard(CHAPTERS_MARKER)
    return sorted(found)


class TemplateRegistry:
    def __init__(self, workspace: WorkspacePaths):
        self.workspace = workspace

    # -- doc-type templates ------------------------------------------------

    def list_doc_types(self) -> list[DocTypeTemplate]:
        out = []
        for f in sorted(self.workspace.doc_type_templates.glob("*.md")):
            out.append(parse_doc_type_template(f.read_text(encoding="utf-8")))
        return out

    def get_doc_type(self, doc_type: str) -> DocTypeTemplate:
        path = self.workspace.doc_type_templates / f"{doc_type}.md"
        if not path.exists():
            raise FileNotFoundError(f"Unknown doc type: {doc_type}")
        return parse_doc_type_template(path.read_text(encoding="utf-8"))

    def save_doc_type(self, tpl: DocTypeTemplate) -> None:
        path = self.workspace.doc_type_templates / f"{tpl.doc_type}.md"
        path.write_text(render_doc_type_template(tpl), encoding="utf-8")

    # -- word templates ------------------------------------------------

    def list_word_templates(self) -> list[str]:
        return sorted(p.name for p in self.workspace.word_templates.glob("*.docx"))

    def save_word_template(self, filename: str, data: bytes) -> str:
        safe_name = Path(filename).name
        dest = self.workspace.word_templates / safe_name
        dest.write_bytes(data)
        return safe_name

    def word_template_path(self, filename: str) -> Path:
        return self.workspace.word_templates / Path(filename).name

    def attach_word_template(self, doc_type: str, filename: str, data: bytes) -> DocTypeTemplate:
        """Save the file into word_templates/ and record it on the doc-type
        template, so export defaults to it without the user picking one.
        """
        saved_name = self.save_word_template(filename, data)
        tpl = self.get_doc_type(doc_type)
        tpl.word_template = saved_name
        self.save_doc_type(tpl)
        return tpl

    def get_template_variables(self, doc_type: str) -> list[str]:
        """Distinct {VARIABLE} tokens in the doc type's attached word
        template. Scanned live (not cached) so it always reflects whatever
        file is currently attached.
        """
        tpl = self.get_doc_type(doc_type)
        if not tpl.word_template:
            return []
        return scan_template_variables(self.word_template_path(tpl.word_template))
