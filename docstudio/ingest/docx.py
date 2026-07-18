from pathlib import Path

from docx import Document as DocxDocument


def extract_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    lines = [f"# {path.name}\n"]
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name if p.style else "") or ""
        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue
        header = [c.text.strip() for c in rows[0].cells]
        lines.append("")
        lines.append("| " + " | ".join(header) + " |")
        lines.append("|" + "---|" * len(header))
        for row in rows[1:]:
            lines.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")

    return "\n\n".join(lines)
