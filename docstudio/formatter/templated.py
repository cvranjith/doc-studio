"""Real DocFormatter: builds the final .docx directly from the document
type's attached Word template, following a convention discovered on a real
corporate template supplied for this project:

- A body paragraph containing exactly ``{CHAPTERS}`` marks where drafted
  chapter content is inserted, styled with the template's own named
  Heading 1-9 / Normal / table styles (so the output looks native to the
  template, not bolted on).
- ``{VARIABLE}`` tokens (ALL_CAPS convention) anywhere in the document —
  body, tables, nested tables, headers, footers — are filled in from
  ``DocManifest.variables`` plus a couple of system-computed values.
- An optional ``##STYLES_START##`` / ``##STYLES_END##`` block (a
  style-reference cheat sheet for whoever built the template) is stripped
  from the delivered document if present.
- ``![alt](assets/x.png)`` image references (pasted into the WYSIWYG chapter
  editor, or one day LLM-generated diagrams) are embedded as real inline
  pictures, scaled down if wider than a normal page margin.

The reasoning engine stays mocked; this formatter does not — it's a real,
working implementation of the DocFormatter protocol (see
docstudio/formatter/__init__.py).
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

import yaml
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from docstudio.models import BuildOptions, DocManifest
from docstudio.store.documents import split_frontmatter
from docstudio.store.templates import VARIABLE_PATTERN, _iter_containers, _iter_paragraphs

STYLES_START_MARKER = "##STYLES_START##"
STYLES_END_MARKER = "##STYLES_END##"
CHAPTERS_MARKER = "{CHAPTERS}"

# CREATION_ON/DOC_NAME are computed at export time, never asked of the user
# (see docstudio/web/js/*: SYSTEM_TEMPLATE_VARIABLES mirrors this list).
SYSTEM_VARIABLES = {"CREATION_ON", "DOC_NAME"}

_YEAR_TOKEN_RE = re.compile(r"\{\s*yyyy\s*\}")
_TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$")
_IMAGE_INLINE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_MAX_IMAGE_WIDTH = Inches(6)

# Inline **bold**/__bold__, *italic*/_italic_, `code` — the WYSIWYG editor's
# turndown.js conversion emits exactly this syntax, so it must round-trip
# into real Word formatting, not literal asterisks. Double-marker
# alternatives are listed first so "**bold**" isn't parsed as two stray
# single-asterisk italics. Known limitation: unlike CommonMark, this doesn't
# apply intraword-underscore rules, so "my_variable_name" would misfire —
# acceptable for prototype-quality chapter prose.
_INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*[^*\s][^*]*?\*|_[^_\s][^_]*?_|`[^`]+?`)")


def _apply_inline_markdown(paragraph, text: str) -> None:
    pos = 0
    for m in _INLINE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**") or token.startswith("__"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(token[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


class TemplatedDocFormatter:
    def build(self, document_dir: Path, word_template: Path, options: BuildOptions) -> Path:
        manifest_data = yaml.safe_load((document_dir / "doc.yaml").read_text(encoding="utf-8"))
        manifest = DocManifest.model_validate(manifest_data)

        document = DocxDocument(str(word_template))

        self._substitute_variables(document, self._resolve_variables(manifest))

        table_style = self._demo_table_style(document)
        anchor = self._find_chapters_anchor(document)
        self._insert_chapters(document, anchor, manifest, document_dir, table_style)
        self._strip_styles_section(document)

        if options.force_draft_watermark:
            self._add_watermark(document)

        self._set_update_fields_on_open(document)

        exports_dir = document_dir.parent.parent / "exports"
        exports_dir.mkdir(parents=True, exist_ok=True)
        suffix = "-draft" if options.force_draft_watermark else ""
        out_path = exports_dir / f"{manifest.slug}-v{manifest.current_version}{suffix}.docx"
        document.save(str(out_path))
        return out_path

    # -- {VARIABLE} substitution -------------------------------------------

    def _resolve_variables(self, manifest: DocManifest) -> dict[str, str]:
        values = dict(manifest.variables)
        values.setdefault("CREATION_ON", manifest.created.strftime("%Y-%m-%d"))
        values.setdefault("DOC_NAME", f"{manifest.slug}-v{manifest.current_version}")
        return values

    def _substitute_variables(self, document, values: dict[str, str]) -> None:
        year = str(date.today().year)
        for container in _iter_containers(document):
            for p in _iter_paragraphs(container):
                self._substitute_paragraph(p, values, year)

    def _substitute_paragraph(self, paragraph, values: dict[str, str], year: str) -> None:
        text = paragraph.text
        if "{" not in text:
            return

        def repl(m: re.Match) -> str:
            key = m.group(1)
            if key == "CHAPTERS":
                return m.group(0)  # handled separately, never treated as a plain variable
            return values.get(key, "")

        new_text = VARIABLE_PATTERN.sub(repl, text)
        new_text = _YEAR_TOKEN_RE.sub(year, new_text)
        if new_text == text:
            return
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

    # -- {CHAPTERS} insertion -----------------------------------------------

    def _find_chapters_anchor(self, document):
        for p in document.paragraphs:
            if p.text.strip() == CHAPTERS_MARKER:
                return p
        return None

    def _insert_chapters(self, document, anchor, manifest: DocManifest, document_dir: Path, table_style: str | None) -> None:
        chapters_dir = document_dir / "chapters"
        for ref in manifest.chapters:
            chapter_path = chapters_dir / ref.file
            if not chapter_path.exists():
                continue
            _, body = split_frontmatter(chapter_path.read_text(encoding="utf-8"))
            for kind, content in _markdown_blocks(body):
                if kind == "table":
                    self._insert_table(document, anchor, content, table_style)
                elif kind == "image":
                    self._insert_image(document, anchor, content, document_dir)
                else:
                    self._insert_paragraph(document, anchor, kind, content)
        if anchor is not None:
            anchor._p.getparent().remove(anchor._p)

    def _insert_paragraph(self, document, anchor, kind: str, text: str):
        p = anchor.insert_paragraph_before("") if anchor is not None else document.add_paragraph()
        style_name = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3", "bullet": "List Bullet"}.get(kind)
        if style_name:
            try:
                p.style = style_name
            except KeyError:
                pass
        _apply_inline_markdown(p, text)
        if kind == "quote":
            for run in p.runs:
                run.italic = True
        return p

    def _insert_image(self, document, anchor, image_info: dict, document_dir: Path):
        """Chapter images (pasted into the WYSIWYG editor, or one day
        LLM-generated diagrams) are stored as markdown ``![alt](assets/x.png)``
        referencing the document's assets/ folder — embedded here as real
        inline pictures, scaled down if wider than a normal page margin.
        """
        image_path = self._resolve_asset_path(document_dir, image_info["path"])
        if image_path is None or not image_path.exists():
            return
        p = anchor.insert_paragraph_before("") if anchor is not None else document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            picture = run.add_picture(str(image_path))
        except Exception:
            return
        if picture.width > _MAX_IMAGE_WIDTH:
            ratio = _MAX_IMAGE_WIDTH / picture.width
            picture.height = int(picture.height * ratio)
            picture.width = _MAX_IMAGE_WIDTH

    def _resolve_asset_path(self, document_dir: Path, rel_path: str) -> Path | None:
        candidate = (document_dir / rel_path).resolve()
        try:
            candidate.relative_to(document_dir.resolve())
        except ValueError:
            return None
        return candidate

    def _insert_table(self, document, anchor, rows: list[list[str]], table_style: str | None):
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        table = document.add_table(rows=len(rows), cols=ncols)
        if table_style:
            try:
                table.style = table_style
            except KeyError:
                pass
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                cell_text = row[ci] if ci < len(row) else ""
                _apply_inline_markdown(table.rows[ri].cells[ci].paragraphs[0], cell_text)
        if anchor is not None:
            anchor._p.addprevious(table._tbl)

    # -- ##STYLES_START##...##STYLES_END## ----------------------------------

    def _demo_table_style(self, document) -> str | None:
        """Table style used by the demo table inside the styles cheat-sheet
        block, if any — applied to tables we generate from chapter markdown.
        """
        in_block = False
        for child in document.element.body:
            tag = child.tag.split("}")[-1]
            if tag == "p":
                text = self._paragraph_text(child)
                if text == STYLES_START_MARKER:
                    in_block = True
                    continue
                if text == STYLES_END_MARKER:
                    break
            if in_block and tag == "tbl":
                for t in document.tables:
                    if t._tbl is child:
                        return t.style.name if t.style else None
        return None

    def _strip_styles_section(self, document) -> None:
        body = document.element.body
        to_remove = []
        removing = False
        for child in list(body):
            tag = child.tag.split("}")[-1]
            text = self._paragraph_text(child) if tag == "p" else None
            if not removing and tag == "p" and text == STYLES_START_MARKER:
                removing = True
                to_remove.append(child)
                continue
            if removing:
                to_remove.append(child)
                if tag == "p" and text == STYLES_END_MARKER:
                    removing = False
                continue
        for el in to_remove:
            body.remove(el)

    def _paragraph_text(self, p_element) -> str:
        return "".join(t.text or "" for t in p_element.findall(f".//{_WORD_NS}t")).strip()

    # -- draft watermark / TOC refresh ---------------------------------------

    def _add_watermark(self, document) -> None:
        paragraphs = document.paragraphs
        wm = paragraphs[0].insert_paragraph_before("") if paragraphs else document.add_paragraph()
        run = wm.add_run("DRAFT — NOT FOR DISTRIBUTION")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)
        wm.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _set_update_fields_on_open(self, document) -> None:
        """Word will recompute the (otherwise stale) Table of Contents page
        numbers the moment the file is opened.
        """
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        document.settings.element.append(el)


def _markdown_blocks(body: str):
    """Yield (kind, content) blocks from chapter markdown. kind is one of
    h1/h2/h3/bullet/quote/para/image/table; table content is a list of row
    lists, image content is {"alt", "path"}.

    Images are detected anywhere within a line, not just lines that are
    nothing but ``![alt](path)`` — the WYSIWYG editor can produce an image
    immediately after typed text with no newline in between (e.g. inserted
    right after a bullet's text), and that still needs to embed as a real
    picture rather than showing up as literal ``![]()`` syntax.
    """
    lines = body.splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [
                [c.strip() for c in tl.strip("|").split("|")]
                for tl in table_lines
                if not _TABLE_SEP_RE.match(tl)
            ]
            if rows:
                yield ("table", rows)
            continue
        if line.startswith("### "):
            kind, text = "h3", line[4:].strip()
        elif line.startswith("## "):
            kind, text = "h2", line[3:].strip()
        elif line.startswith("# "):
            kind, text = "h1", line[2:].strip()
        elif line.startswith("- ") or line.startswith("* "):
            kind, text = "bullet", line[2:].strip()
        elif line.startswith(">"):
            kind, text = "quote", line.lstrip("> ").strip()
        else:
            kind, text = "para", line
        yield from _split_inline_images(kind, text)
        i += 1


def _split_inline_images(kind: str, text: str):
    pos = 0
    any_image = False
    for m in _IMAGE_INLINE_RE.finditer(text):
        any_image = True
        before = text[pos : m.start()].strip()
        if before:
            yield (kind, before)
        yield ("image", {"alt": m.group(1), "path": m.group(2)})
        pos = m.end()
    if not any_image:
        yield (kind, text)
    else:
        after = text[pos:].strip()
        if after:
            yield (kind, after)
