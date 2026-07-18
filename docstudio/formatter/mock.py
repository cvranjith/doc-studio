"""MockDocFormatter — produces a plausible stub .docx (title page + raw
chapter text via python-docx). Validation of chapter status / open questions
lives in docstudio/api/export.py, which passes the operator's choice down via
BuildOptions.force_draft_watermark.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from docstudio.models import BuildOptions, DocManifest
from docstudio.store.documents import split_frontmatter


class MockDocFormatter:
    def build(self, document_dir: Path, word_template: Path, options: BuildOptions) -> Path:
        manifest_data = yaml.safe_load((document_dir / "doc.yaml").read_text(encoding="utf-8"))
        manifest = DocManifest.model_validate(manifest_data)

        doc = Document()
        self._add_title_page(doc, manifest, word_template, options)

        for ref in manifest.chapters:
            chapter_path = document_dir / "chapters" / ref.file
            if not chapter_path.exists():
                continue
            _, body = split_frontmatter(chapter_path.read_text(encoding="utf-8"))
            self._add_chapter(doc, body)

        exports_dir = document_dir.parent.parent / "exports"
        exports_dir.mkdir(parents=True, exist_ok=True)
        suffix = "-draft" if options.force_draft_watermark else ""
        out_path = exports_dir / f"{manifest.slug}-v{manifest.current_version}{suffix}.docx"
        doc.save(str(out_path))
        return out_path

    def _add_title_page(self, doc: Document, manifest: DocManifest, word_template: Path, options: BuildOptions) -> None:
        if options.force_draft_watermark:
            watermark = doc.add_paragraph()
            watermark.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = watermark.add_run("DRAFT — NOT FOR DISTRIBUTION")
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(manifest.title)
        run.bold = True
        run.font.size = Pt(28)

        for label, value in [
            ("Client", manifest.client or "—"),
            ("Document Type", manifest.doc_type),
            ("Status", manifest.status),
            ("Version", str(manifest.current_version)),
            ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
            ("Template", word_template.name if word_template else "(mock — no template applied)"),
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"{label}: {value}").font.size = Pt(11)

        doc.add_page_break()

    def _add_chapter(self, doc: Document, body: str) -> None:
        for line in body.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith(">"):
                p = doc.add_paragraph()
                p.add_run(stripped.lstrip("> ").strip()).italic = True
            elif stripped.startswith("|"):
                continue  # tables rendered as-is-skipped in the mock formatter
            else:
                doc.add_paragraph(stripped)
        doc.add_page_break()
